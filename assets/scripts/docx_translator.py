# -*- coding: utf-8 -*-
import os
import sys
import json
import re
import argparse
import shutil
from docx import Document
from deep_translator import GoogleTranslator

# Regex nhận diện tiếng Trung
CHINESE_REGEX = re.compile(r'[\u4e00-\u9fff]')

def contains_chinese(text):
    if not isinstance(text, str):
        return False
    return bool(CHINESE_REGEX.search(text))

def translate_text_online(text, engine, target_lang="vi", api_key=None):
    if not text or not text.strip():
        return text
    try:
        if engine == "gemini" and api_key:
            import requests
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            prompt = f"Bạn là một biên dịch viên chuyên nghiệp. Hãy dịch đoạn văn bản sau đây sang tiếng Việt. Chỉ trả về bản dịch, không giải thích gì thêm:\n\n{text}"
            data = {
                "contents": [{
                    "parts": [{"text": prompt}]
                }]
            }
            response = requests.post(url, headers=headers, json=data, timeout=10)
            if response.status_code == 200:
                result = response.json()
                translated = result['candidates'][0]['content']['parts'][0]['text'].strip()
                if translated.startswith('"') and translated.endswith('"'):
                    translated = translated[1:-1]
                return translated
            else:
                print(f"[WARN] Gemini API failed with status {response.status_code}, fallback to Google Free.", file=sys.stderr)
        
        return GoogleTranslator(source='auto', target=target_lang).translate(text)
    except Exception as e:
        print(f"[WARN] Translation error: {str(e)}", file=sys.stderr)
        return text

def translate_runs(runs, dict_data, engine, target_lang, api_key):
    """
    Dịch văn bản ở cấp độ Runs để giữ nguyên định dạng của từng Run (in đậm, in nghiêng, màu sắc...)
    """
    total_translated = 0
    
    # Chuẩn hóa từ điển
    normalized_dict = {str(k).strip().lower(): str(v).strip() for k, v in dict_data.items() if k}
    
    for run in runs:
        text = run.text
        if text and contains_chinese(text):
            val_str = text.strip()
            val_lower = val_str.lower()
            
            # Lớp 1: Khớp từ điển hoàn toàn
            if val_lower in normalized_dict:
                run.text = text.replace(val_str, normalized_dict[val_lower])
                total_translated += 1
                continue
                
            # Dịch một phần bằng từ điển
            matched_dict = False
            for key_word, trans_word in normalized_dict.items():
                if key_word in val_lower:
                    pattern = re.compile(re.escape(key_word), re.IGNORECASE)
                    val_str = pattern.sub(trans_word, val_str)
                    matched_dict = True
            
            if matched_dict:
                run.text = text.replace(run.text.strip(), val_str)
                total_translated += 1
                if not contains_chinese(val_str):
                    continue
            
            # Lớp 2: Dịch online
            translated = translate_text_online(val_str, engine, target_lang, api_key)
            if translated and translated != val_str:
                run.text = text.replace(run.text.strip(), translated)
                total_translated += 1
                
    return total_translated

def process_docx(src_path, dest_path, dict_data, engine, target_lang, api_key):
    # Tạo bản sao file gốc để ghi đè giữ nguyên định dạng
    shutil.copy2(src_path, dest_path)
    
    doc = Document(dest_path)
    total_translated = 0
    
    # 1. Dịch trong các Paragraphs thông thường
    for para in doc.paragraphs:
        if para.runs:
            total_translated += translate_runs(para.runs, dict_data, engine, target_lang, api_key)
            
    # 2. Dịch trong các Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Dịch các paragraph trong ô bảng
                for para in cell.paragraphs:
                    if para.runs:
                        total_translated += translate_runs(para.runs, dict_data, engine, target_lang, api_key)
                        
    # 3. Quét kiểm tra chất lượng (Verification)
    remaining_chinese_runs = []
    # Quét lại toàn bộ xem còn chữ Trung không
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.text and contains_chinese(run.text):
                remaining_chinese_runs.append({
                    "location": f"Para {i}",
                    "value": run.text
                })
                
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    for run in para.runs:
                        if run.text and contains_chinese(run.text):
                            remaining_chinese_runs.append({
                                "location": f"Table {t_idx}, Row {r_idx}, Cell {c_idx}",
                                "value": run.text
                            })
                            
    doc.save(dest_path)
    
    return {
        "success": True,
        "total_translated": total_translated,
        "remaining_chinese_count": len(remaining_chinese_runs),
        "remaining_cells": remaining_chinese_runs[:20]
    }

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="JA Word Document Translator Tool")
    parser.add_argument("--src", required=True, help="Path to source Word file")
    parser.add_argument("--dest", required=True, help="Path to destination Word file")
    parser.add_argument("--dict", default="{}", help="JSON string of translation dictionary")
    parser.add_argument("--engine", default="google", help="Translation engine (google, gemini)")
    parser.add_argument("--lang", default="vi", help="Target language code")
    parser.add_argument("--key", default="", help="API Key for translation engine")
    
    args = parser.parse_args()
    
    try:
        try:
            dict_data = json.loads(args.dict)
        except Exception:
            dict_data = {}
            
        result = process_docx(
            src_path=args.src,
            dest_path=args.dest,
            dict_data=dict_data,
            engine=args.engine,
            target_lang=args.lang,
            api_key=args.key
        )
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        error_res = {
            "success": False,
            "error": str(e)
        }
        print(json.dumps(error_res, ensure_ascii=False))
        sys.exit(1)
